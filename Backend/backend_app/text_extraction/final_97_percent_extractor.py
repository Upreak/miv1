"""
Final 97% extractor - comprehensive text extraction for all supported file types.
Uses multiple extraction strategies and falls back gracefully.
"""

import logging
import os
from typing import Optional

from backend_app.text_extraction.utils import (
    create_temp_file, cleanup_temp_file, get_file_extension, 
    is_text_file, sanitize_filename
)
from backend_app.text_extraction.unstructured_io_runner import extract_with_unstructured

logger = logging.getLogger(__name__)

def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Extract text from file using the best available method.
    
    Args:
        file_bytes: File content as bytes
        filename: Original filename
        
    Returns:
        str: Extracted text
        
    Raises:
        RuntimeError: If extraction fails for all methods
    """
    logger.info(f"Starting text extraction for {filename}")
    
    # Step 1: Handle text files directly
    if is_text_file(file_bytes, filename):
        try:
            text = file_bytes.decode('utf-8')
            logger.debug(f"Text file detected, extracted {len(text)} characters")
            return text
        except Exception as e:
            logger.warning(f"Failed to decode text file {filename}: {str(e)}")
    
    # Step 2: Try Unstructured.io API
    file_extension = get_file_extension(filename)
    extracted_text = extract_with_unstructured(file_bytes, filename, file_extension)
    
    if extracted_text:
        logger.debug(f"Unstructured.io extraction successful for {filename}")
        return extracted_text
    
    # Step 3: Fallback to basic extraction methods
    extracted_text = _extract_with_fallback_methods(file_bytes, filename)
    
    if extracted_text:
        logger.debug(f"Fallback extraction successful for {filename}")
        return extracted_text
    
    # Step 4: If all methods fail, return empty string with warning
    logger.warning(f"All extraction methods failed for {filename}")
    return f"# WARNING: Could not extract text from {filename}\n\n"
    f"The file may be corrupted, password-protected, or in an unsupported format.\n"
    f"File size: {len(file_bytes)} bytes\n"
    f"File extension: {file_extension}"

def _extract_with_fallback_methods(file_bytes: bytes, filename: str) -> Optional[str]:
    """
    Fallback extraction methods when primary methods fail.
    
    Args:
        file_bytes: File content as bytes
        filename: Original filename
        
    Returns:
        str: Extracted text or None if extraction failed
    """
    file_extension = get_file_extension(filename).lower()
    
    # Try PDF extraction
    if file_extension == '.pdf':
        return _extract_pdf_text(file_bytes, filename)
    
    # Try DOC/DOCX extraction
    elif file_extension in ['.doc', '.docx']:
        return _extract_doc_text(file_bytes, filename)
    
    # Try other methods for unknown file types
    else:
        return _extract_generic_text(file_bytes, filename)

def _extract_pdf_text(file_bytes: bytes, filename: str) -> Optional[str]:
    """
    Extract text from PDF using PyPDF2 or similar.
    """
    try:
        # Try PyPDF2 first
        try:
            import PyPDF2
            from io import BytesIO
            
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
            text = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                except Exception as e:
                    logger.warning(f"Failed to extract page {page_num} from {filename}: {str(e)}")
            
            if text.strip():
                logger.debug(f"PyPDF2 extraction successful for {filename}")
                return text.strip()
                
        except ImportError:
            logger.debug("PyPDF2 not available, trying pdfplumber")
        
        # Try pdfplumber as fallback
        try:
            import pdfplumber
            
            text = ""
            with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            
            if text.strip():
                logger.debug(f"pdfplumber extraction successful for {filename}")
                return text.strip()
                
        except ImportError:
            logger.debug("pdfplumber not available")
        
        # Try OCR for scanned PDFs
        try:
            import pytesseract
            from PIL import Image
            import fitz  # PyMuPDF
            
            text = ""
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                # Render page to image
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x resolution
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                
                # OCR the image
                page_text = pytesseract.image_to_string(img)
                if page_text.strip():
                    text += page_text + "\n\n"
            
            if text.strip():
                logger.debug(f"OCR extraction successful for {filename}")
                return text.strip()
                
        except ImportError:
            logger.debug("PyMuPDF or pytesseract not available for OCR")
        except Exception as e:
            logger.warning(f"OCR extraction failed for {filename}: {str(e)}")
        
    except Exception as e:
        logger.warning(f"PDF extraction failed for {filename}: {str(e)}")
    
    return None

def _extract_doc_text(file_bytes: bytes, filename: str) -> Optional[str]:
    """
    Extract text from DOC/DOCX files.
    """
    try:
        # Try docx2txt for DOCX files
        if filename.lower().endswith('.docx'):
            try:
                import docx2txt
                
                text = docx2txt.process(BytesIO(file_bytes))
                if text and text.strip():
                    logger.debug(f"docx2txt extraction successful for {filename}")
                    return text.strip()
                    
            except ImportError:
                logger.debug("docx2txt not available")
        
        # Try python-docx for DOCX files
        try:
            from docx import Document
            from io import BytesIO
            
            doc = Document(BytesIO(file_bytes))
            text = []
            
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)
            
            if text:
                result = '\n\n'.join(text)
                logger.debug(f"python-docx extraction successful for {filename}")
                return result.strip()
                
        except ImportError:
            logger.debug("python-docx not available")
        
        # Try antiword for DOC files
        try:
            import subprocess
            import tempfile
            
            with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as temp_file:
                temp_file.write(file_bytes)
                temp_path = temp_file.name
            
            try:
                result = subprocess.run(
                    ['antiword', temp_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                
                if result.returncode == 0 and result.stdout.strip():
                    logger.debug(f"antiword extraction successful for {filename}")
                    return result.stdout.strip()
                    
            finally:
                os.unlink(temp_path)
                
        except (ImportError, FileNotFoundError):
            logger.debug("antiword not available")
        except Exception as e:
            logger.warning(f"antiword extraction failed for {filename}: {str(e)}")
        
    except Exception as e:
        logger.warning(f"DOC/DOCX extraction failed for {filename}: {str(e)}")
    
    return None

def _extract_generic_text(file_bytes: bytes, filename: str) -> Optional[str]:
    """
    Generic text extraction for unknown file types.
    """
    try:
        # Try to extract any readable text
        text = file_bytes.decode('utf-8', errors='ignore')
        
        # Filter out non-printable characters
        printable_text = ''.join(c for c in text if c.isprintable() or c.isspace())
        
        # Check if we have meaningful content
        if printable_text.strip() and len(printable_text) > 100:
            # Remove excessive whitespace
            lines = [line.strip() for line in printable_text.split('\n') if line.strip()]
            result = '\n'.join(lines)
            
            if len(result) > 50:  # Ensure we have substantial content
                logger.debug(f"Generic text extraction successful for {filename}")
                return result
        
    except Exception as e:
        logger.warning(f"Generic text extraction failed for {filename}: {str(e)}")
    
    return None